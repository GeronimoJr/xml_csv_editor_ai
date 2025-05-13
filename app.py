import streamlit as st
import requests
import tempfile
import os
import re
import traceback
import json
from datetime import datetime
from pydrive2.auth import GoogleAuth
from pydrive2.drive import GoogleDrive
from oauth2client.service_account import ServiceAccountCredentials
import pandas as pd
import numpy as np
import io
from docx import Document
import tiktoken
import math
import xml.etree.ElementTree as ET

st.set_page_config(page_title="Tłumacz plików AI", layout="centered")
st.title("AI Tłumacz plików CSV, XML, Excel i Word")
st.markdown("""
To narzędzie umożliwia tłumaczenie zawartości plików CSV, XML, XLS, XLSX, DOC i DOCX za pomocą wybranego modelu LLM.
Prześlij plik, wybierz język docelowy oraz model.
""")

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    user = st.text_input("Login")
    password = st.text_input("Hasło", type="password")
    if st.button("Zaloguj"):
        if user == st.secrets.get("APP_USER") and password == st.secrets.get("APP_PASSWORD"):
            st.session_state.authenticated = True
        else:
            st.error("Nieprawidłowy login lub hasło")
    st.stop()

if "translated_text" not in st.session_state:
    st.session_state.translated_text = None
if "output_bytes" not in st.session_state:
    st.session_state.output_bytes = None

drive_folder_id = st.secrets.get("GOOGLE_DRIVE_FOLDER_ID")
service_account_json = st.secrets.get("GOOGLE_DRIVE_CREDENTIALS_JSON")

uploaded_file = st.file_uploader("Wgraj plik do przetłumaczenia", type=["xml", "csv", "xls", "xlsx", "doc", "docx"])
target_lang = st.selectbox("Język docelowy", ["en", "pl", "de", "fr", "es", "it"])
model = st.selectbox("Wybierz model LLM (OpenRouter)", [
    "openai/gpt-4o-mini",
    "openai/gpt-4o",
    "openai/gpt-4-turbo",
    "anthropic/claude-3-opus",
    "mistralai/mistral-7b-instruct",
    "google/gemini-pro"
])
api_key = st.secrets["OPENROUTER_API_KEY"]

MODEL_PRICES = {
    "openai/gpt-4o-mini": {"prompt": 0.15, "completion": 0.6},
    "mistralai/mistral-7b-instruct": {"prompt": 0.2, "completion": 0.2},
    "google/gemini-pro": {"prompt": 0.25, "completion": 0.5},
}

def extract_xml_texts_and_paths(elem, path=""):
    texts = []
    if elem.text and elem.text.strip():
        texts.append((f"{path}/text", elem.text.strip()))
    for k, v in elem.attrib.items():
        texts.append((f"{path}/@{k}", v))
    for i, child in enumerate(elem):
        child_path = f"{path}/{child.tag}[{i}]"
        texts.extend(extract_xml_texts_and_paths(child, child_path))
    return texts

def insert_translations_into_xml(elem, translations, path=""):
    if elem.text and elem.text.strip():
        key = f"{path}/text"
        if key in translations:
            elem.text = translations[key]
    for k in elem.attrib:
        key = f"{path}/@{k}"
        if key in translations:
            elem.attrib[k] = translations[key]
    for i, child in enumerate(elem):
        child_path = f"{path}/{child.tag}[{i}]"
        insert_translations_into_xml(child, translations, child_path)

if uploaded_file:
    file_type = uploaded_file.name.split(".")[-1].lower()
    raw_bytes = uploaded_file.read()
    try:
        if file_type == "xml":
            encoding_declared = re.search(br'<\?xml[^>]*encoding=["\']([^"\']+)["\']', raw_bytes)
            encodings_to_try = [encoding_declared.group(1).decode('ascii')] if encoding_declared else []
            encodings_to_try += ["utf-8", "iso-8859-2", "windows-1250", "utf-16"]
            for enc in encodings_to_try:
                try:
                    file_contents = raw_bytes.decode(enc)
                    tree = ET.ElementTree(ET.fromstring(file_contents))
                    root = tree.getroot()
                    break
                except Exception:
                    continue
            else:
                st.error("Nie udało się odczytać pliku – nieznane kodowanie lub błąd składni XML.")
                st.stop()
            pairs = extract_xml_texts_and_paths(root)
            keys, lines = zip(*pairs) if pairs else ([], [])
        elif file_type == "csv":
            for enc in ["utf-8", "iso-8859-2", "windows-1250", "utf-16"]:
                try:
                    df = pd.read_csv(io.BytesIO(raw_bytes), encoding=enc)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                st.error("Nie udało się odczytać pliku CSV – nieznane kodowanie.")
                st.stop()
            lines = []
            cell_indices = []
            for row_idx, row in df.iterrows():
                for col_idx, cell in enumerate(row):
                    lines.append(str(cell))
                    cell_indices.append((row_idx, df.columns[col_idx]))
        elif file_type in ["xls", "xlsx"]:
            df = pd.read_excel(io.BytesIO(raw_bytes))
            lines = []
            cell_indices = []
            for row_idx, row in df.iterrows():
                for col_idx, cell in enumerate(row):
                    lines.append(str(cell))
                    cell_indices.append((row_idx, df.columns[col_idx]))
        elif file_type in ["doc", "docx"]:
            doc = Document(io.BytesIO(raw_bytes))
            lines = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            lines.append(cell.text.strip())
        else:
            st.error("Nieobsługiwany typ pliku.")
            st.stop()

        enc = tiktoken.encoding_for_model("gpt-4")
        chunk_size = 10000
        chunks, current_chunk, current_tokens = [], [], 0
        for i, line in enumerate(lines):
            token_len = len(enc.encode(line))
            if current_tokens + token_len > chunk_size:
                chunks.append(current_chunk)
                current_chunk, current_tokens = [], 0
            current_chunk.append((i, line))
            current_tokens += token_len
        if current_chunk:
            chunks.append(current_chunk)

        prompt_tokens = sum(len(enc.encode(l)) for _, l in sum(chunks, []))
        completion_tokens = int(prompt_tokens * 1.2)
        total_tokens = prompt_tokens + completion_tokens

        pricing = MODEL_PRICES.get(model, {"prompt": 1.0, "completion": 1.0})
        cost_prompt = prompt_tokens / 1_000_000 * pricing["prompt"]
        cost_completion = completion_tokens / 1_000_000 * pricing["completion"]
        cost_total = cost_prompt + cost_completion

        st.info(f"Szacunkowe zużycie tokenów: ~{prompt_tokens} (prompt) + ~{completion_tokens} (output) = ~{total_tokens} tokenów, w {len(chunks)} częściach")
        st.info(f"Szacunkowy koszt tłumaczenia: ~${cost_total:.4f} USD")

        if st.button("Przetłumacz plik"):
            translated_pairs = []
            for i, chunk in enumerate(chunks):
                with st.spinner(f"Tłumaczenie części {i + 1} z {len(chunks)}..."):
                    content = "\n".join(l for _, l in chunk)
                    prompt = f"Przetłumacz na język {target_lang}. Zwróć każdą linię w oryginalnej kolejności, bez numeracji.\n\n{content}"
                    res = requests.post("https://openrouter.ai/api/v1/chat/completions",
                        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                        json={"model": model, "messages": [
                            {"role": "system", "content": "Tłumacz precyzyjnie bez zmiany formatu."},
                            {"role": "user", "content": prompt}
                        ]})
                    result = res.json()["choices"][0]["message"]["content"].splitlines()
                    for (idx, _), translated in zip(chunk, result):
                        translated_pairs.append((idx, translated.strip()))
            translated_pairs.sort()

            with tempfile.TemporaryDirectory() as tmpdir:
                output_path = os.path.join(tmpdir, f"output.{file_type}")
                if file_type == "xml":
                    translated_map = dict(translated_pairs)
                    insert_translations_into_xml(root, translated_map)
                    tree.write(output_path, encoding="utf-8", xml_declaration=True)
                elif file_type in ["csv", "xls", "xlsx"]:
                    translated_df = df.copy()
                    for (idx, (row_idx, col_name)) in enumerate(cell_indices):
                        translated_df.at[row_idx, col_name] = translated_pairs[idx][1]
                    if file_type == "csv":
                        translated_df.to_csv(output_path, index=False)
                    else:
                        translated_df.to_excel(output_path, index=False)
                elif file_type in ["doc", "docx"]:
                    new_doc = Document()
                    index = 0
                    for p in doc.paragraphs:
                        if p.text.strip():
                            new_doc.add_paragraph(translated_pairs[index][1])
                            index += 1
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    cell.text = translated_pairs[index][1]
                                    index += 1
                    new_doc.save(output_path)

                with open(output_path, "rb") as f:
                    st.session_state.output_bytes = f.read()

                if drive_folder_id and service_account_json:
                    creds_dict = json.loads(service_account_json)
                    scope = ["https://www.googleapis.com/auth/drive"]
                    credentials = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
                    gauth = GoogleAuth()
                    gauth.credentials = credentials
                    drive = GoogleDrive(gauth)

                    now = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
                    result_filename = f"translated_output_{now}.{file_type}"
                    result_file = drive.CreateFile({"title": result_filename, "parents": [{"id": drive_folder_id}]})
                    result_file.SetContentFile(output_path)
                    result_file.Upload()
                    st.success("Plik zapisany na Twoim Google Drive ✅")

                st.success("Tłumaczenie zakończone. Plik gotowy do pobrania.")

    except Exception as e:
        st.error("Błąd podczas przetwarzania:")
        st.exception(traceback.format_exc())

if st.session_state.output_bytes:
    st.download_button("📁 Pobierz przetłumaczony plik", data=st.session_state.output_bytes, file_name=f"translated_output.{file_type}", mime="application/octet-stream")
